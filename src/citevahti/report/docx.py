"""Word (.docx) bridge — the world researchers actually live in.

Export the Citation-Integrity Report to .docx, and import a .docx manuscript back to
Markdown so it can enter the normal paste → review flow. Backed by ``python-docx``,
which is a CORE dependency so Word features work out of the box — including in the
no-terminal ``.mcpb``. The guard below is only a dead-man's switch for a broken install.
"""

from __future__ import annotations

from io import BytesIO

from ..schemas.report import STATE_CODE, ClaimReport

_STATE_TITLE = {
    "accepted": "Accepted", "needs_support": "Needs support",
    "review_needed": "Review needed", "decision_recorded": "Decision recorded",
    "untestable": "Untestable (out of indexed scope)",
}
_ORDER = ("needs_support", "review_needed", "decision_recorded", "accepted", "untestable")


def _require_docx():
    try:
        import docx  # noqa: F401  (python-docx)
    except ImportError as exc:
        raise RuntimeError(
            "Word (.docx) support needs python-docx, which ships with CiteVahti. "
            "Your install looks incomplete — reinstall CiteVahti (or, from source, "
            "pip install python-docx).") from exc
    from docx import Document
    return Document


def render_docx(report: ClaimReport, *, title: str = "Citation-Integrity Report") -> bytes:
    """The report as a .docx byte stream — same content as the Markdown/HTML exports."""
    Document = _require_docx()
    c = report.counts
    d = Document()
    d.add_heading(title, level=0)
    d.add_paragraph(f"Generated {report.generated_at} · {report.total} claim(s) tested.")

    table = d.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"      # built-in; guard in case a stripped template lacks it
    except Exception:  # noqa: BLE001
        pass
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "State", "Count"
    for k, lbl in (("accepted", "Accepted (supported citation)"), ("needs_support", "Needs support"),
                   ("review_needed", "Review needed"), ("decision_recorded", "Decision recorded"),
                   ("untestable", "Untestable (out of indexed scope)")):
        row = table.add_row().cells
        row[0].text, row[1].text = lbl, str(c.get(k, 0))

    for state in _ORDER:
        rows = [r for r in report.rows if r.state == state]
        if not rows:
            continue
        d.add_heading(_STATE_TITLE[state], level=1)
        for r in rows:
            code = STATE_CODE[r.state].strip()
            d.add_heading(f"[{code}] {r.claim_text}", level=2)
            if r.manuscript_location:
                d.add_paragraph(f"Location: {r.manuscript_location}")
            if r.proposed_revision:
                by = r.proposed_revision_by or "human"
                d.add_paragraph(f"Pending revision ({by}-proposed): {r.proposed_revision}")
            for e in r.evidence:
                ids = " · ".join(p for p in (f"PMID {e.pmid}" if e.pmid else "",
                                                  f"DOI {e.doi}" if e.doi else "") if p)
                ai = "hidden (blinded)" if e.ai_support == "hidden" else (e.ai_support or "—")
                dec = f" · decision: {e.final_decision}" if e.final_decision else ""
                flag = "  [RETRACTED]" if e.retracted else ""
                line = f"{e.title or '(untitled)'}{flag}" + (f" — {ids}" if ids else "")
                line += f"  (human: {e.human_support or '—'} · AI: {ai}{dec})"
                d.add_paragraph(line, style="List Bullet")
            if not r.evidence:
                d.add_paragraph("no candidate evidence linked yet", style="List Bullet")

    d.add_paragraph("")
    foot = d.add_paragraph()
    foot.add_run("Scope & limitations. ").bold = True
    foot.add_run("States record citation support from the blinded human → AI → "
                 "adjudication workflow — not clinical or scientific truth. Coverage is "
                 "the claims the author entered into the ledger; the audit chain is "
                 "tamper-evident, not forgery-proof.")

    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def docx_to_markdown(data: bytes) -> str:
    """A .docx manuscript as Markdown (headings by style, paragraphs as text), so it can
    enter the normal paste → claim-extraction flow. Lossy on tables/figures by design —
    claims are prose."""
    Document = _require_docx()
    doc = Document(BytesIO(data))
    out: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        style = (getattr(p.style, "name", "") or "").lower()
        if not text:
            out.append("")
            continue
        if style == "title" or "heading 1" in style:
            out.append("# " + text)
        elif "heading 2" in style:
            out.append("## " + text)
        elif "heading 3" in style:
            out.append("### " + text)
        elif "heading" in style:
            out.append("#### " + text)
        else:
            out.append(text)
    # collapse 3+ blank lines, ensure a trailing newline
    md = "\n".join(out).strip()
    while "\n\n\n" in md:
        md = md.replace("\n\n\n", "\n\n")
    return md + "\n"
